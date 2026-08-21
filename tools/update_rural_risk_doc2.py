# -*- coding: utf-8 -*-
"""将备份后页面修改同步到《智能风险研判（农村自建房）风险清单20260816.docx》，修改内容用蓝色字体显示"""
import docx
from docx.shared import RGBColor

SRC = r'e:\风险管控0618\智能风险研判（农村自建房）风险清单20260816.docx'
DST = SRC
BLUE = RGBColor(0x00, 0x70, 0xC0)

doc = docx.Document(SRC)
changes = []

def set_para_blue_text(p, new_text):
    base = p.runs[0] if p.runs else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    run.font.color.rgb = BLUE
    if base is not None:
        try:
            run.font.name = base.font.name
            run.font.size = base.font.size
            run.bold = base.bold
        except Exception:
            pass

def cell_write(cell, new_text, tag):
    set_para_blue_text(cell.paragraphs[0], new_text)
    # 清除多余段落
    for p in cell.paragraphs[1:]:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    changes.append(tag)

def cell_replace(cell, old, new, tag):
    for p in cell.paragraphs:
        if old in p.text:
            set_para_blue_text(p, p.text.replace(old, new))
            changes.append(tag)
            return True
    return False

# ---------- TABLE 0 输入字段表 ----------
t0 = doc.tables[0]
for row in t0.rows:
    name = row.cells[0].text.strip()
    if name == '分项':
        cell_write(row.cells[1], '下拉选择框＋自定义输入', 'T0 分项控件')
        cell_write(row.cells[3],
            '选项随风险主项联动：建筑结构信息→地基基础、墙体、梁、柱、楼、屋盖、次要构件；'
            '建筑基本信息→建成时间、设计方式、改扩建情况；选择"自定义..."后展开输入框手动录入',
            'T0 分项说明')
    elif name == '分项内容':
        cell_replace(row.cells[3], '支持输入1～100个字符', '支持输入1～100个字符', 'T0 分项内容')
    elif name == '风险识别':
        cell_write(row.cells[1], '富文本编辑器', 'T0 风险识别控件')
        cell_write(row.cells[3],
            '带工具栏的富文本编辑器：支持加粗、斜体、下划线、删除线、有序/无序列表排版，'
            '支持插入图片、文档、视频附件（单个附件最大20MB），底部实时显示已插入附件数量',
            'T0 风险识别说明')
    elif name == '可能导致事故':
        cell_write(row.cells[1], '标签多选＋自定义添加', 'T0 事故控件')
        cell_write(row.cells[2], '是', 'T0 事故必填')
        cell_write(row.cells[3],
            '预设事故标签：坍塌、高处坠落、触电、火灾，点击标签多选；'
            '点击"+ 自定义"按钮展开输入框，输入后回车或失焦添加为自定义标签；'
            '实时显示"已选N项"', 'T0 事故说明')

# ---------- 段落 13 必填校验规则 ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('13. 必填校验规则'):
        set_para_blue_text(p,
            '13. 必填校验规则：基础信息步骤（第1步）的风险主项、风险等级、分项、分项内容、风险识别、'
            '可能导致事故均为必填项，标签以红色"*"标识；未填写或仅包含空格时不得保存并提示；'
            '已关联预防措施为必填项；选择"条件启用"时启用条件类型和条件值为必填。')
        changes.append('规则13 必填校验')
    elif t.startswith('4. 新增交互'):
        set_para_blue_text(p,
            '4. 新增交互：点击"新增风险"打开新增分步向导弹窗，表单重置，风险主项默认"建筑结构信息"、'
            '分项选项随主项联动、风险等级默认"第二类（黄）"、监管状态默认"启用"；'
            '通过"下一步／上一步"在3个步骤间切换，步骤标签同步高亮；'
            '弹窗底部不提供"取消"按钮，仅可通过右上角关闭按钮关闭；'
            '第3步保存按钮文案为"确认新增"。')
        changes.append('交互4 新增交互')
    elif t.startswith('5. 编辑交互'):
        set_para_blue_text(p,
            '5. 编辑交互：点击"编辑"打开编辑分步向导弹窗并回填全部字段，序号只读；'
            '分项按原值回填，非预设分项以"自定义"形式回填；可能导致事故按标签选中态回填，'
            '非预设事故自动追加为自定义标签；保存后更新数据并提示"保存成功"。'
            '保存期间按钮置灰，防止重复提交；弹窗底部不提供"取消"按钮。')
        changes.append('交互5 编辑交互')

# ---------- TABLE 11 提示文案 ----------
t11 = doc.tables[11]
for row in t11.rows:
    scene = row.cells[0].text.strip()
    if scene == '可能导致事故未选择':
        pass
# 追加一行：可能导致事故未选择
new_row = t11.add_row()
cell_write(new_row.cells[0], '可能导致事故未选择', 'T11 新增行-场景')
cell_write(new_row.cells[1], '请填写可能导致事故', 'T11 新增行-文案')

# ---------- TABLE 13 埋点 ----------
t13 = doc.tables[13]
for row in t13.rows:
    ev = row.cells[0].text.strip()
    if ev == '新增风险条目提交':
        cell_replace(row.cells[1], '点击保存', '点击"确认新增"', 'T13 新增埋点触发时机')

# ---------- TABLE 15 新增验收 ----------
t15 = doc.tables[15]
rows_to_remove = []
for row in t15.rows:
    scene = row.cells[0].text.strip()
    if scene == '【正常】新增条目':
        cell_replace(row.cells[1], '点击"保存"', '点击"确认新增"', 'T15 新增条目-操作')
        if '确认新增' not in row.cells[1].text:
            cell_replace(row.cells[1], '点击“保存”', '点击"确认新增"', 'T15 新增条目-操作2')
    elif scene == '【正常】第3步按钮':
        cell_write(row.cells[2], '不展示"下一步"按钮，展示"确认新增"按钮', 'T15 第3步按钮')
    elif scene == '【正常】取消新增':
        cell_write(row.cells[0], '【正常】关闭新增弹窗', 'T15 取消新增-场景')
        cell_write(row.cells[1], '填写部分内容后点击右上角关闭按钮', 'T15 取消新增-操作')
        cell_write(row.cells[2], '弹窗关闭，系统不新增条目，列表和统计数据不变；弹窗底部无"取消"按钮', 'T15 取消新增-预期')
    elif scene == '【异常】分项为空':
        cell_write(row.cells[1], '不填写分项，完成其他必填项后点击"确认新增"', 'T15 分项为空-操作')
        cell_write(row.cells[2], '系统不保存数据，并提示"请填写分项"', 'T15 分项为空-预期')
    elif scene == '【异常】分项内容为空':
        cell_write(row.cells[1], '不填写分项内容后点击"确认新增"', 'T15 分项内容为空-操作')
    elif scene == '【异常】风险识别为空':
        cell_write(row.cells[1], '不在风险识别富文本编辑器中输入内容后点击"确认新增"', 'T15 风险识别为空-操作')
    elif scene == '【异常】重复点击保存':
        cell_replace(row.cells[1], '点击两次"保存"', '点击两次"确认新增"', 'T15 重复点击-操作')
        cell_replace(row.cells[1], '点击两次“保存”', '点击两次"确认新增"', 'T15 重复点击-操作2')
    elif scene == '【异常】新增接口失败':
        cell_replace(row.cells[1], '点击"保存"', '点击"确认新增"', 'T15 接口失败-操作')
        cell_replace(row.cells[1], '点击“保存”', '点击"确认新增"', 'T15 接口失败-操作2')
# 追加事故校验用例行
r = t15.add_row()
cell_write(r.cells[0], '【异常】可能导致事故未选择', 'T15 新增-事故场景')
cell_write(r.cells[1], '不选择任何事故标签，完成其他必填项后点击"确认新增"', 'T15 新增-事故操作')
cell_write(r.cells[2], '系统不保存数据，并提示"请填写可能导致事故"', 'T15 新增-事故预期')

# ---------- TABLE 16 编辑验收 ----------
t16 = doc.tables[16]
for row in t16.rows:
    scene = row.cells[0].text.strip()
    if scene == '【正常】打开编辑弹窗':
        cell_write(row.cells[2],
            '弹窗回填该条目原有主项、等级、分项（含自定义分项）、分项内容、风险识别富文本内容、'
            '可能导致事故标签选中态、预防措施、工作依据、备注及监管状态，序号只读；底部无"取消"按钮',
            'T16 打开编辑弹窗-预期')

doc.save(DST)
print('已保存:', DST)
for c in changes:
    print(' -', c)
print('共', len(changes), '处修改')
