# -*- coding: utf-8 -*-
"""调整《智能风险研判（农村自建房）风险清单20260816.docx》，修改内容用红色文字显示"""
import docx
from docx.shared import RGBColor

SRC = r'e:\风险管控0618\智能风险研判（农村自建房）风险清单20260816.docx'
DST = r'e:\风险管控0618\智能风险研判（农村自建房）风险清单20260816.docx'
RED = RGBColor(0xFF, 0x00, 0x00)

doc = docx.Document(SRC)
changes = []

def set_para_red_text(p, new_text):
    """清空段落所有 run，写入红色新文本（保留首个 run 的字体样式）"""
    base = p.runs[0] if p.runs else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    run.font.color.rgb = RED
    if base is not None:
        try:
            run.font.name = base.font.name
            run.font.size = base.font.size
            run.bold = base.bold
        except Exception:
            pass

def replace_in_para(p, old, new, tag):
    if old in p.text:
        set_para_red_text(p, p.text.replace(old, new))
        changes.append(tag)
        return True
    return False

def cell_para_replace(cell, old, new, tag):
    done = False
    for p in cell.paragraphs:
        if old in p.text:
            set_para_red_text(p, p.text.replace(old, new))
            done = True
    if done:
        changes.append(tag)
    return done

# ---------- 1. 段落修改 ----------
# 【业务规则】16 看板统计规则
for p in doc.paragraphs:
    t = p.text
    if t.startswith('16. 看板统计规则'):
        set_para_red_text(p,
            '16. 看板统计规则：监管状态看板支持“按街镇”“按房屋类型”两个维度统计各分组下启用、条件启用、禁用条目数量；'
            '按街镇维度支持从镇/街道下钻至行政村，下钻后可通过面包屑返回“全部街镇”；'
            '支持在分组卡片上点击星标手动设置或取消“重点关注”，被关注卡片高亮显示“重点关注”标识；'
            '初始默认将未启用率≥50%的分组自动推荐为重点关注；重点关注设置保存在浏览器本地（localStorage），刷新后保留。')
        changes.append('业务规则16 看板统计规则')
    elif t.startswith('11. 看板交互'):
        set_para_red_text(p,
            '11. 看板交互：点击“按街镇／按房屋类型”切换统计维度；按街镇维度下点击乡镇卡片下钻至行政村层级，'
            '点击面包屑“全部街镇”返回乡镇层级；点击卡片右上角星标设置或取消该分组的重点关注，星标点击不触发下钻；'
            '点击“查看全部”跳转完整状态变更记录列表。')
        changes.append('交互逻辑11 看板交互')

# ---------- 2. 表格修改 ----------
# TABLE 7 看板输出字段
t7 = doc.tables[7]
for row in t7.rows:
    c0 = row.cells[0].text.strip()
    if c0 == '高亮显示重点关注':
        cell_para_replace(row.cells[1],
            '人工设置重点关注，卡片高亮提示关注',
            '卡片右上角提供星标按钮，人工点击设置/取消重点关注；被关注卡片高亮并展示“重点关注”标识；'
            '初始默认未启用率≥50%的分组自动标记为重点关注；设置结果保存在浏览器本地，刷新后保留',
            'TABLE7 高亮显示重点关注')

# TABLE 13 埋点：看板维度切换
t13 = doc.tables[13]
for row in t13.rows:
    cell_para_replace(row.cells[1], '按地域／按房屋类型', '按街镇／按房屋类型', 'TABLE13 埋点-看板维度切换')

# TABLE 14 验收（1）
t14 = doc.tables[14]
for row in t14.rows:
    scene = row.cells[0].text.strip()
    if scene == '【正常】看板下钻':
        cell_para_replace(row.cells[1], '按地域维度下', '按街镇维度下', 'TABLE14 看板下钻-操作')
    elif scene == '【正常】看板返回':
        cell_para_replace(row.cells[1], '点击返回按钮', '点击面包屑“全部街镇”', 'TABLE14 看板返回-操作')
        cell_para_replace(row.cells[2], '恢复为镇/街道层级统计', '恢复为全部街镇层级统计', 'TABLE14 看板返回-预期')
    elif scene == '【正常】未启用率高亮':
        cell_para_replace(row.cells[0], '【正常】未启用率高亮', '【正常】星标设置重点关注', 'TABLE14 行标题')
        set_para_red_text(row.cells[1].paragraphs[0], '点击某一分组卡片右上角的星标按钮')
        set_para_red_text(row.cells[2].paragraphs[0],
            '该卡片高亮并展示“重点关注”标识，星标变为选中态；再次点击取消关注，高亮移除；'
            '初始默认未启用率≥50%的分组卡片自动标记为重点关注')
        changes.append('TABLE14 未启用率高亮→星标设置重点关注')

# TABLE 22 审计与埋点验收
t22 = doc.tables[22]
for row in t22.rows:
    if '看板下钻' in row.cells[0].text:
        cell_para_replace(row.cells[1], '按地域维度', '按街镇维度', 'TABLE22 埋点-看板下钻')

doc.save(DST)
print('已保存:', DST)
for c in changes:
    print(' -', c)
