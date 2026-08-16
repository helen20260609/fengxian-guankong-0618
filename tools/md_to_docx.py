# -*- coding: utf-8 -*-
"""将整合后的需求规格说明书md转换为Word文档"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD = r'D:\风险管控\doc\智能风险研判——第三方服务机构功能需求规格说明书.md'
DOCX = r'D:\风险管控\doc\智能风险研判——第三方服务机构功能需求规格说明书.docx'

with open(MD, 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

doc = Document()

# 全局默认字体
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.size = Pt({1: 16, 2: 14, 3: 12, 4: 11}[i])
    hs.font.bold = True

def add_runs(par, text):
    """处理行内 **bold** 标记"""
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = par.add_run(part[2:-2])
            run.bold = True
        elif part:
            par.add_run(part)

def flush_table(tbl_lines):
    rows = []
    for ln in tbl_lines:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        if all(re.fullmatch(r':?-{2,}:?', c or '---') for c in cells):
            continue  # 分隔行
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ''
            par = cell.paragraphs[0]
            text = row[ci] if ci < len(row) else ''
            text = text.replace('**', '')
            run = par.add_run(text)
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
    doc.add_paragraph()

i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    # 表格块
    if s.startswith('|'):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl.append(lines[i])
            i += 1
        flush_table(tbl)
        continue
    if not s:
        i += 1
        continue
    if s == '---':
        doc.add_paragraph()
        i += 1
        continue
    m = re.match(r'^(#{1,4})\s+(.*)$', s)
    if m:
        level = len(m.group(1))
        doc.add_heading(m.group(2).replace('**', ''), level=level)
        i += 1
        continue
    par = doc.add_paragraph()
    add_runs(par, s)
    i += 1

# 封面若干行居中（文档前5行）
for idx, par in enumerate(doc.paragraphs[:6]):
    t = par.text.strip()
    if t.startswith('上海市奉贤区政府采购') or t.startswith('——') or \
       t.startswith('上海市奉贤区建管委') or t == '上海联平科技有限公司' or t == '2026年8月15日':
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in par.runs:
            if t.startswith('上海市奉贤区建管委'):
                run.font.size = Pt(18)
                run.bold = True

doc.save(DOCX)
print('OK', DOCX)

