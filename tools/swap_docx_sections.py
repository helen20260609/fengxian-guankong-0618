from docx import Document
from docx.shared import Pt
import copy
import os
import re

DOCX_PATH = r'e:\风险管控0618\智能风险研判（农村自建房）风险清单20260818.docx'
BACKUP_PATH = r'e:\风险管控0618\智能风险研判（农村自建房）风险清单20260818.bak.docx'

# 复制备份
import shutil
shutil.copy2(DOCX_PATH, BACKUP_PATH)
print(f"已创建备份: {BACKUP_PATH}")

doc = Document(DOCX_PATH)

# 找到1.1.1和1.1.2以及下一个同级标题1.1.3的位置
indices = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # 精确匹配标题开头
    if re.match(r'^1\.1\.1\s', text):
        indices['1.1.1'] = i
    elif re.match(r'^1\.1\.2\s', text):
        indices['1.1.2'] = i
    elif re.match(r'^1\.1\.3\s', text):
        indices['1.1.3'] = i
    elif re.match(r'^1\.2\s', text):
        indices['1.2'] = i

print("找到的索引:", indices)

start_111 = indices['1.1.1']
start_112 = indices['1.1.2']
end_112 = indices.get('1.1.3', indices.get('1.2'))

# 如果没有1.1.3，则取1.2；否则取1.1.3
if end_112 is None:
    raise ValueError("未找到1.1.3或1.2标题，无法确定1.1.2章节结束位置")

# 提取段落对象（注意：这里提取的是Element对象，需要在body层面重排）
body = doc.element.body

# 提取1.1.1段落（从start_111到start_112-1）
sec111 = []
for idx in range(start_111, start_112):
    sec111.append(doc.paragraphs[idx]._element)

# 提取1.1.2段落（从start_112到end_112-1）
sec112 = []
for idx in range(start_112, end_112):
    sec112.append(doc.paragraphs[idx]._element)

# 在body中移除这些段落元素
for el in sec111 + sec112:
    body.remove(el)

# 按互换后的顺序重新插入到原1.1.1的位置（即body中start_111对应的位置）
# body中start_111之前的元素数量
insert_before = start_111
# 注意body中可能还有表格等元素，用paragraph索引定位不一定精确
# 更安全：在body中找到原来start_111前面的元素作为锚点
# 但因为已经移除了sec111和sec112，body中start_111位置的元素现在是什么？
# 原来的start_111前面有insert_before个段落，但body中可能还包含表格

# 使用paragraph的element作为锚点：找到当前body中原来在start_111位置之前的段落
# 由于移除了sec111, sec112，原来start_112后面的段落现在前移了
# 我们需要把新的1.1.2插入到原来1.1.1的位置

# 找到锚点：body中当前在insert_before位置的段落（即原来1.1.2结束后的下一个元素）
# 由于body不仅仅包含paragraph，这里用paragraph计数
anchor_para_idx = start_111
if anchor_para_idx < len(doc.paragraphs):
    anchor = doc.paragraphs[anchor_para_idx]._element
else:
    anchor = None

# 先插入1.1.2的内容
if anchor is not None:
    for el in sec112:
        anchor.addprevious(el)
else:
    for el in sec112:
        body.append(el)

# 再插入1.1.1的内容到1.1.2之前（即在当前1.1.2内容之前）
# 由于sec112已经插入到anchor之前，现在anchor仍然是原来的锚点
for el in sec111:
    anchor.addprevious(el)

# 修改标题编号：互换后内容对应的标题需要调整
# 找到互换后的1.1.1和1.1.2标题元素并交换它们的编号文本
# 注意：元素已经互换位置，但标题文字没变，所以需要更新
# 新的第一个章节原本是1.1.2，现在应该改为1.1.1
# 新的第二个章节原本是1.1.1，现在应该改为1.1.2

# 获取刚插入的元素（按顺序）
# 由于addprevious会改变元素位置，最后插入的在最前面，所以顺序是 sec111 后 sec112
# 但实际上我们插入顺序是：先sec112到anchor前，再sec111到anchor前
# 所以最终body中从原位置开始依次是：sec111的元素、sec112的元素、anchor
# 标题编号需要交换

def set_title_number(element, new_number):
    """把标题中的 '1.1.x' 替换为 new_number"""
    # element 是 CT_P，text在多个run中
    # 简单处理：遍历所有w:t节点，替换第一个匹配
    from docx.oxml.ns import qn
    found = False
    for t in element.iter(qn('w:t')):
        if not found and re.search(r'1\.1\.\d+', t.text):
            t.text = re.sub(r'1\.1\.\d+', new_number, t.text, count=1)
            found = True

# 找到sec111和sec112的第一个段落（应该是标题）
title_111_new = sec111[0]  # 这个标题原本是1.1.1，现在位置变成1.1.2
title_112_new = sec112[0]  # 这个标题原本是1.1.2，现在位置变成1.1.1

set_title_number(title_112_new, '1.1.1')
set_title_number(title_111_new, '1.1.2')

doc.save(DOCX_PATH)
print("已保存修改")
